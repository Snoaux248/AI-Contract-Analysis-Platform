from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Dict

import fitz
from docx import Document


@dataclass
class ParsedDocument:
    text: str
    metadata: Dict[str, str]


class DocumentParser:
    """Parse supported contract file formats into plain text."""

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}

    def parse(self, file_path: Path) -> ParsedDocument:
        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {suffix}")

        if suffix == ".txt":
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".pdf":
            text = self._parse_pdf(file_path)
        else:
            text = self._parse_docx(file_path)

        normalized_text = self._normalize_text(text)
        if not normalized_text.strip():
            raise ValueError("Document is empty or unreadable")

        return ParsedDocument(
            text=normalized_text,
            metadata={
                "filename": file_path.name,
                "file_type": suffix.replace(".", ""),
            },
        )

    def _parse_pdf(self, file_path: Path) -> str:
        doc = fitz.open(file_path)
        pages = [page.get_text("text") for page in doc]
        doc.close()
        return "\n".join(pages)

    def _parse_docx(self, file_path: Path) -> str:
        doc = Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text)

    def _normalize_text(self, text: str) -> str:
        text = text.replace("\xa0", " ")
        text = re.sub(r"\r\n?", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
